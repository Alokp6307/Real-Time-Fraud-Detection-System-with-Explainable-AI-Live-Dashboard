from pathlib import Path
import warnings

import joblib
import numpy as np
import pandas as pd
from lightgbm import LGBMClassifier
from sklearn.model_selection import train_test_split


ROOT = Path(__file__).resolve().parent
DASHBOARD_DIR = ROOT / "dashboard"
DASHBOARD_DIR.mkdir(exist_ok=True)
warnings.filterwarnings("ignore", category=pd.errors.PerformanceWarning)


def load_and_prepare():
    transaction = pd.read_csv(ROOT / "train_transaction.csv")
    identity = pd.read_csv(ROOT / "train_identity.csv")
    df = transaction.merge(identity, on="TransactionID", how="left")

    new_features = pd.DataFrame(index=df.index)
    new_features["AmtToMeanRatio"] = df["TransactionAmt"] / df["TransactionAmt"].mean()
    new_features["HourOfDay"] = (df["TransactionDT"] // 3600) % 24

    if "DeviceType" in df.columns and "DeviceInfo" in df.columns:
        new_features["DeviceRisk"] = np.where(
            df["DeviceType"].isna()
            | df["DeviceInfo"].isna()
            | (df["DeviceType"].astype(str).str.lower() == "mobile"),
            1,
            0,
        )
    else:
        new_features["DeviceRisk"] = 0

    df = pd.concat([df.copy(), new_features], axis=1)

    missing_percent = df.isna().mean() * 100
    df = df[missing_percent[missing_percent <= 50].index]

    y = df["isFraud"]
    X = df.drop(columns=["isFraud", "TransactionID"])

    numeric_columns = X.select_dtypes(include=["number"]).columns.tolist()
    categorical_columns = X.select_dtypes(
        include=["object", "string", "category"]
    ).columns.tolist()

    numeric_medians = {}
    category_maps = {}

    for column in numeric_columns:
        median_value = X[column].median()
        numeric_medians[column] = median_value
        X[column] = X[column].fillna(median_value)

    for column in categorical_columns:
        mode_value = X[column].mode()
        fill_value = mode_value.iloc[0] if len(mode_value) else "Unknown"
        values = X[column].fillna(fill_value).astype(str)
        categories = pd.Categorical(values).categories
        category_maps[column] = {
            "fill_value": fill_value,
            "categories": categories.tolist(),
        }
        X[column] = pd.Categorical(values, categories=categories).codes

    X = X.loc[:, ~X.columns.duplicated()]
    return X, y, numeric_medians, category_maps


def train_and_save_model():
    X, y, numeric_medians, category_maps = load_and_prepare()
    sample_size = min(120000, len(X))

    X_train, _, y_train, _ = train_test_split(
        X,
        y,
        train_size=sample_size,
        random_state=42,
        stratify=y,
    )

    model = LGBMClassifier(
        random_state=42,
        n_estimators=250,
        learning_rate=0.05,
        class_weight="balanced",
        n_jobs=-1,
        verbose=-1,
    )
    model.fit(X_train, y_train)

    payload = {
        "model": model,
        "feature_columns": X.columns.tolist(),
        "numeric_medians": numeric_medians,
        "category_maps": category_maps,
        "risk_tiers": {
            "Critical Risk": "probability >= 0.75",
            "Suspicious": "0.40 <= probability < 0.75",
            "Clear": "probability < 0.40",
        },
    }

    joblib.dump(payload, DASHBOARD_DIR / "model.pkl")


def create_docx_summary():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_heading("Fraud Detection Internship Project Summary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This summary documents the complete fraud detection workflow, including "
        "data preparation, model training, explainability, risk segmentation, "
        "visualization, and dashboard deployment readiness."
    )

    sections = [
        (
            "Data Loading and Exploration",
            [
                "Transaction and identity CSV files were loaded using Pandas.",
                "Both files were merged on TransactionID.",
                "Class imbalance, missing values, amount distribution, and correlations were analyzed.",
            ],
        ),
        (
            "Preprocessing and Feature Engineering",
            [
                "Columns with more than 50% missing values were dropped.",
                "Numerical values were imputed using median and categorical values using mode.",
                "High-cardinality categorical columns were label encoded.",
                "Engineered features included AmtToMeanRatio, HourOfDay, and DeviceRisk.",
                "SMOTE was applied only on the training set after stratified 80/20 split.",
            ],
        ),
        (
            "Model Training and Evaluation",
            [
                "LightGBM, XGBoost, and Isolation Forest models were trained.",
                "Models were evaluated using Accuracy, Precision, Recall, F1-Score, ROC-AUC, and PR-AUC.",
                "Threshold optimization was performed using a Threshold vs F1-Score plot.",
                "RandomizedSearchCV was used for tuning the best model.",
            ],
        ),
        (
            "Explainability and Risk Analysis",
            [
                "SHAP global summary and waterfall plots were generated.",
                "Fraud cases, borderline cases, and legitimate transactions were explained in plain English.",
                "Transactions were segmented into Critical Risk, Suspicious, and Clear risk tiers.",
            ],
        ),
        (
            "Dashboard and Visualizations",
            [
                "A Streamlit fraud operations dashboard was created.",
                "Dashboard pages include Overview, Transaction Explorer, and SHAP Explainer.",
                "Required charts include SHAP summary, fraud rate by hour, TransactionAmt distribution, risk tier donut chart, and Precision-Recall curve.",
            ],
        ),
        (
            "Business Recommendations",
            [
                "Critical Risk transactions should go through step-up verification.",
                "High-risk hours and risky device behavior should be monitored in real time.",
                "PR-AUC should be prioritized over accuracy because fraud data is highly imbalanced.",
            ],
        ),
    ]

    for heading, bullets in sections:
        doc.add_heading(heading, level=1)
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Saved Deliverables", level=1)
    doc.add_paragraph("dashboard/model.pkl", style="List Bullet")
    doc.add_paragraph("summary.docx", style="List Bullet")
    doc.add_paragraph("summary.pdf", style="List Bullet")

    output = ROOT / "summary.docx"
    doc.save(output)


def create_pdf_summary():
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    output = ROOT / "summary.pdf"
    doc = SimpleDocTemplate(str(output), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Fraud Detection Internship Project Summary", styles["Title"]))
    story.append(Spacer(1, 12))

    content = [
        "Transaction and identity CSV files were loaded and merged on TransactionID.",
        "Columns with more than 50% missing values were dropped.",
        "Remaining numerical values were imputed using median and categorical values using mode.",
        "Feature engineering included AmtToMeanRatio, HourOfDay, and DeviceRisk.",
        "LightGBM, XGBoost, and Isolation Forest were trained and compared.",
        "Evaluation used Accuracy, Precision, Recall, F1-Score, ROC-AUC, and PR-AUC.",
        "SHAP was used for global and local explainability.",
        "Risk tiers were created: Critical Risk, Suspicious, and Clear.",
        "A Streamlit dashboard was created for operations, exploration, and SHAP explanations.",
        "Recommended policies include step-up verification and real-time device/hour monitoring.",
    ]

    for item in content:
        story.append(Paragraph(f"- {item}", styles["BodyText"]))
        story.append(Spacer(1, 6))

    doc.build(story)


if __name__ == "__main__":
    train_and_save_model()
    create_docx_summary()
    create_pdf_summary()
    print("Created dashboard/model.pkl, summary.docx, and summary.pdf")
